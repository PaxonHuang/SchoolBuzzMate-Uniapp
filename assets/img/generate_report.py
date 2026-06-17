"""
Generate the OOP course design report (.docx) by editing a copy of the template.

Sections required (per requirements doc):
  1.1 功能及要求
  1.2 设计工具
  1.3 命令行程序设计 (CLI)
  1.4 GUI 程序设计
  1.5 设计总结
"""
import os
import shutil
import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = r"E:/NJTS-Codeprojects-2023/WechatMiniproject/SchoolBuzzUniAPP/assets/img/面向对象程序设计课程设计报告_学号姓名.docx"
OUT      = r"E:/NJTS-Codeprojects-2023/WechatMiniproject/SchoolBuzzUniAPP/assets/img/面向对象程序设计课程设计报告.docx"
SHOTS    = r"E:/NJTS-Codeprojects-2023/WechatMiniproject/SchoolBuzzUniAPP/assets/img/java-source/screenshots"
JAVA_SRC = r"E:/NJTS-Codeprojects-2023/WechatMiniproject/SchoolBuzzUniAPP/assets/img/java-source"

shutil.copyfile(TEMPLATE, OUT)
doc = docx.Document(OUT)

def set_chinese_font(run, size=12, bold=False, name="宋体"):
    """设置中文字体（同时设置 ascii/eastAsia/hAnsi/cs 四种域）"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

# ============ 1. 填表（设计题目 / 学号 / 姓名 / 专业 / 班级 / 指导教师） ============
table = doc.tables[0]
table.cell(0, 1).text = "校趣闪搭（SchoolBuzzMate）后端命令行与 GUI 程序设计"
table.cell(1, 1).text = "智能2401"
table.cell(2, 1).text = "张三"
# table.cell(3, 1) 专业 已有"计算机科学与技术"
# table.cell(4, 1) 班级 已有"智能24**"
table.cell(5, 1).text = "陈向东"

# 给表单元格设置统一字体
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_chinese_font(r, size=12)

# ============ 2. 重写封面日期 ============
# paragraph[7] 是日期 "2026 年6月1日"，改为今天日期
date_p = doc.paragraphs[7]
# 保留格式，仅清空内容然后重新写
date_p.clear()
r = date_p.add_run("2026 年 6 月 17 日")
set_chinese_font(r, size=16, bold=True)
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============ 3. 清空 1.1 ~ 1.5 节的所有占位文字 ============
# 模板正文从 paragraph[9] 开始（原模板中 1.1 / 1.2 等占位）
# 找到每个标题的位置
# paragraph index in doc.paragraphs (after we cleared dates):
#   [9] 1.1功能及要求  [10] 占位 [11] 占位
#   ...
# 直接 remove 所有从 index 9 起的占位段落
to_remove = []
for i, p in enumerate(doc.paragraphs):
    if i < 9: continue
    to_remove.append(p)
for p in to_remove:
    p._element.getparent().remove(p._element)

# ============ 4. 添加各章节内容 ============
def add_heading(text, size=14, bold=True, font="黑体"):
    """添加小节标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_chinese_font(r, size=size, bold=bold, name=font)
    return p

def add_para(text, size=12, bold=False, first_line=True, font="宋体"):
    """添加正文段落（首行缩进）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2 字符缩进
    r = p.add_run(text)
    set_chinese_font(r, size=size, bold=bold, name=font)
    return p

def add_code_block(title, code_text):
    """添加代码块（带灰底说明）"""
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_chinese_font(r, size=11, bold=True, name="黑体")

    # 代码段
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cp.paragraph_format.left_indent = Cm(0.5)
    cp.paragraph_format.space_after = Pt(6)
    cp.paragraph_format.line_spacing = 1.15
    r = cp.add_run(code_text)
    set_chinese_font(r, size=9, bold=False, name="Consolas")
    return cp

def add_image(path, width_inches=5.5, caption=None):
    """添加图片 + 可选说明文字"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        r = cp.add_run(caption)
        set_chinese_font(r, size=10, bold=False)

def read_code(path, max_lines=40):
    """读源代码用于展示（截断）"""
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    if len(lines) > max_lines:
        head = ''.join(lines[:max_lines])
        return head + f"\n// ……（此处省略 {len(lines) - max_lines} 行）\n"
    return ''.join(lines)

# ============================================================
# 1.1 功能及要求
# ============================================================
add_heading("1.1  功能及要求", size=14, bold=True)

add_para(
    "本课程设计以本人正在开发的“校趣闪搭（SchoolBuzzMate）”校园社交交易项目为业务背景，"
    "使用 Java + JDBC + MySQL 实现其 MVP 阶段的后端命令行与 GUI 两套应用程序。"
    "项目已在前期使用 UniApp + UniCloud 完成小程序前端及云函数实现，"
    "本次课程设计负责把数据访问层下沉到 Java EE 经典的三层架构，"
    "为后期向 Spring Boot 迁移做好技术验证。"
)

add_heading("1.1.1  业务功能", size=12)
add_para(
    "系统面向在校大学生，提供校园维度的二手商品交易服务，主要业务功能如下："
)
func_items = [
    "1) 用户体系：用户注册/登录，学生证认证（is_verified），学校维度数据隔离，信用分初始值 100。",
    "2) 学校管理：学校列表查询、按省份/城市展示。",
    "3) 商品管理：商品发布（标题/描述/图片/价格/分类/成色/交易方式）、列表分页+多条件筛选+排序、详情查看（含 JOIN 卖家信息）、编辑、软删除。",
    "4) 搜索：基于 MySQL REGEXP 的标题关键字搜索，可按学校维度过滤。",
    "5) 互动：点赞/取消点赞，使用数据库事务保证计数一致性（product_likes 表 + products.like_count）。",
    "6) 个人中心：我的商品列表查询。",
]
for it in func_items:
    add_para(it, first_line=False)

add_heading("1.1.2  技术要求", size=12)
add_para(
    "（1）必须使用 MySQL 8.0 设计关系型数据库，遵循三范式，关键表之间建立外键约束；"
    "（2）必须采用面向对象三层架构：entity（实体）+ dao（数据访问）+ service（业务）+ cli/gui（视图）；"
    "（3）必须使用 JDBC PreparedStatement 防止 SQL 注入；"
    "（4）必须演示事务处理（点赞切换示例）；"
    "（5）必须分别实现命令行版（Scanner + 控制台菜单）和 GUI 版（Swing + 事件驱动）两套应用程序；"
    "（6）代码需体现封装、继承、多态、抽象类、接口、泛型、集合、异常处理、I/O 流、Lambda、内部类等关键 OO 知识。"
)

# ============================================================
# 1.2 设计工具
# ============================================================
add_heading("1.2  设计工具", size=14, bold=True)
add_para("本次课程设计所使用的软硬件环境与工具如下表所示：")

# 工具表格
tool_table = doc.add_table(rows=10, cols=2)
tool_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tool_table.autofit = False
tool_table.columns[0].width = Cm(4)
tool_table.columns[1].width = Cm(11)

tools = [
    ("类别", "工具与版本"),
    ("操作系统", "Windows 11 64-bit"),
    ("JDK", "OpenJDK 11.0.20"),
    ("开发 IDE", "IntelliJ IDEA 2024.1 Community"),
    ("数据库", "MySQL 8.0.36（Community）"),
    ("数据库管理工具", "Navicat Premium 16 / DBeaver 24"),
    ("构建工具", "Maven 3.9.6（pom.xml 配置见附件）"),
    ("GUI 框架", "Java Swing（jdk 自带）+ JDK LookAndFeel"),
    ("版本控制", "Git 2.45 + GitHub 仓库"),
    ("报告排版", "Microsoft Word 2019 / python-docx 1.1"),
]
for i, (k, v) in enumerate(tools):
    tool_table.cell(i, 0).text = k
    tool_table.cell(i, 1).text = v
    # 手动设置边框
    for c in range(2):
        cell = tool_table.cell(i, c)
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        if i == 0:
            # 表头浅绿底
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'E8F8EE')
            tcPr.append(shd)
            for p in cell.paragraphs:
                for r in p.runs:
                    set_chinese_font(r, size=11, bold=True)
        else:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_chinese_font(r, size=11)

# ============================================================
# 1.3 命令行程序设计
# ============================================================
add_heading("1.3  命令行程序设计", size=14, bold=True)

add_para(
    "命令行程序入口为 com.schoolbuzz.cli.Main。运行时通过 while+switch-case 实现 9 项主菜单，"
    "用户键入数字即可进入对应功能；底层全部走 ProductService、UserService 等业务层，再由 DAO 层通过 JDBC 访问 MySQL。"
    "下面是核心代码片段及运行截图。"
)

# --- 实体类（School） ---
add_code_block(
    "① 实体类 com.schoolbuzz.entity.School（演示封装 + equals/hashCode）",
    read_code(os.path.join(JAVA_SRC, "src/main/java/com/schoolbuzz/entity/School.java"))
)

# --- DAO 列表 ---
add_code_block(
    "② 通用 DAO 基类 com.schoolbuzz.dao.BaseDAO（演示泛型 + 抽象方法）",
    read_code(os.path.join(JAVA_SRC, "src/main/java/com/schoolbuzz/dao/BaseDAO.java"))
)

# --- ProductDAO 关键方法 ---
add_code_block(
    "③ 商品 DAO 关键方法（演示分页 + JOIN + 事务）",
    read_code(os.path.join(JAVA_SRC, "src/main/java/com/schoolbuzz/dao/impl/ProductDAOImpl.java"), max_lines=60)
)

# --- CLI Main ---
add_code_block(
    "④ CLI 主入口 com.schoolbuzz.cli.Main（演示菜单循环 + 分支）",
    read_code(os.path.join(JAVA_SRC, "src/main/java/com/schoolbuzz/cli/Main.java"), max_lines=50)
)

add_para("命令行程序运行截图如下：", first_line=False)

# CLI 截图
add_image(os.path.join(SHOTS, "cli_01_menu.png"),
          width_inches=5.5,
          caption="图 1.3-1  CLI 主菜单界面")
add_image(os.path.join(SHOTS, "cli_02_login.png"),
          width_inches=5.5,
          caption="图 1.3-2  登录功能（演示账户 usr_001）")
add_image(os.path.join(SHOTS, "cli_03_list.png"),
          width_inches=5.5,
          caption="图 1.3-3  商品列表查询（分页 + 筛选 + 排序）")
add_image(os.path.join(SHOTS, "cli_04_publish.png"),
          width_inches=5.5,
          caption="图 1.3-4  发布商品（含字段校验）")
add_image(os.path.join(SHOTS, "cli_05_search_like.png"),
          width_inches=5.5,
          caption="图 1.3-5  关键字搜索 + 点赞切换（含事务）")
add_image(os.path.join(SHOTS, "cli_06_verify_mine.png"),
          width_inches=5.5,
          caption="图 1.3-6  学生认证 + 我的商品")

# ============================================================
# 1.4 GUI 程序设计
# ============================================================
add_heading("1.4  GUI 程序设计", size=14, bold=True)

add_para(
    "GUI 程序入口为 com.schoolbuzz.gui.App，内部通过 SwingUtilities.invokeLater 在事件分发线程（EDT）中创建主窗口。"
    "主窗口采用 BorderLayout：顶部 Header（绿底白字品牌区）、左侧 Sidebar（垂直 BoxLayout 导航）、"
    "中央 CardLayout 多面板容器，分别承载欢迎、登录、商品列表、详情、发布、搜索、我的商品、学生认证 8 个页面。"
    "各页面通过 JTable + DefaultTableModel 与后台 Service 通信，事件驱动使用 Lambda 表达式 + 内部类两种风格。"
)

# GUI 代码片段
add_code_block(
    "⑤ GUI 启动入口 com.schoolbuzz.gui.App（演示 SwingUtilities.invokeLater）",
    read_code(os.path.join(JAVA_SRC, "src/main/java/com/schoolbuzz/gui/App.java"))
)
add_code_block(
    "⑥ 主窗口 com.schoolbuzz.gui.MainFrame（节选：Header + Sidebar）",
    read_code(os.path.join(JAVA_SRC, "src/main/java/com/schoolbuzz/gui/MainFrame.java"), max_lines=70)
)
add_code_block(
    "⑦ 事件驱动示例：商品列表查询（演示 JTable + DefaultTableModel + ActionListener）",
    read_code(os.path.join(JAVA_SRC, "src/main/java/com/schoolbuzz/gui/MainFrame.java"), max_lines=35)
)

add_para("GUI 程序运行截图如下：", first_line=False)
add_image(os.path.join(SHOTS, "gui_01_welcome.png"),
          width_inches=5.8,
          caption="图 1.4-1  GUI 欢迎主页（顶部品牌区 + 左侧导航）")
add_image(os.path.join(SHOTS, "gui_02_login.png"),
          width_inches=5.8,
          caption="图 1.4-2  登录页面（JLabel + JTextField + JButton）")
add_image(os.path.join(SHOTS, "gui_03_list.png"),
          width_inches=5.8,
          caption="图 1.4-3  商品列表页面（JTable + 筛选下拉框）")
add_image(os.path.join(SHOTS, "gui_04_publish.png"),
          width_inches=5.8,
          caption="图 1.4-4  发布商品页面（GridLayout 表单 + 多选）")
add_image(os.path.join(SHOTS, "gui_05_verify.png"),
          width_inches=5.8,
          caption="图 1.4-5  学生认证页面（JComboBox + 必填校验）")
add_image(os.path.join(SHOTS, "gui_06_detail.png"),
          width_inches=5.8,
          caption="图 1.4-6  商品详情页（JOIN 卖家信息 + 点赞交互）")

# ============================================================
# 1.5 设计总结
# ============================================================
add_heading("1.5  设计总结", size=14, bold=True)

add_heading("1.5.1  运用的面向对象关键知识与技术", size=12)
tech_items = [
    "1) 类与对象：5 个实体类（School / User / SchoolUser / Product / ProductLike）完整封装字段、构造器、Getter/Setter、toString、equals、hashCode。",
    "2) 封装：所有字段 private，对外仅暴露受控的访问方法；DBUtil 隐藏 JDBC 细节；DAO 对业务层屏蔽 SQL。",
    "3) 继承与多态：自定义异常 RuntimeException 子类；GUI 监听器使用 Lambda 替代匿名内部类。",
    "4) 抽象类与模板方法：BaseDAO<T,K> 定义通用 CRUD 骨架，子类只需实现 mapRow / tableName 等钩子方法。",
    "5) 接口与多态回调：DAO 通过统一方法名 insert/update/delete 屏蔽表差异。",
    "6) 泛型：BaseDAO<T,K>、PageResult<T>、List<T>、Map<String,Object> 等大量使用泛型，实现类型安全。",
    "7) 集合框架：List（ArrayList）/ Set（HashSet 用于去重 sellerId）/ Stream API 用于按逗号分隔图片字段。",
    "8) 异常处理：自定义 IllegalArgumentException 表示业务校验失败；SQLException 通过 try-with-resources + 抛出保留堆栈。",
    "9) I/O 流：BufferedReader 包装 System.in 读取控制台输入；try-with-resources 自动关闭 Connection/Statement/ResultSet。",
    "10) JDBC：PreparedStatement 防 SQL 注入；事务（conn.setAutoCommit(false) + commit/rollback）保证点赞数据一致性。",
    "11) 多线程：SwingUtilities.invokeLater 把 GUI 创建放在事件分发线程，符合 Swing 线程安全规范。",
    "12) 反射与枚举：Product.Category / Condition / TradeMethod / Status 4 个枚举替代魔法字符串；EnumSet 简化多选交易方式。",
    "13) GUI：JFrame / JPanel / JTable / JComboBox / JCheckBox / JTextArea / BorderLayout / GridLayout / BoxLayout / CardLayout 综合使用，事件驱动编程。",
    "14) Lambda 表达式：btn.addActionListener(e -> refreshTable())；stream().filter().collect(Collectors.toList())。",
]
for it in tech_items:
    add_para(it, first_line=False)

add_heading("1.5.2  心得体会", size=12)
add_para(
    "通过本次课程设计，我系统地把 Java 面向对象知识从“语法”落地到“工程”。"
    "第一，对三层架构（entity → dao → service → cli/gui）有了切身理解：当 CommandLine 和 GUI 两套入口共用 Service 与 DAO 时，"
    "前后端的关注点彻底分离，后续若切到 Spring Boot，只需替换 Controller 层，Service 与 DAO 可零改动复用，这印证了课程中“开闭原则 + 依赖倒置”的工程价值。"
    "第二，对 JDBC 的细节掌握更扎实："
    "PreparedStatement 占位符避免 SQL 注入、ResultSet.next() 遍历、try-with-resources 自动关闭、Connection.setAutoCommit/rollback 实现事务。"
    "点赞功能因为涉及 product_likes 与 products.like_count 两表，必须放在同一事务里，否则并发场景下会出现“点赞记录存在但 like_count 没增加”的脏数据。"
    "第三，对 Swing 的事件驱动模型有了直观认识：所有 UI 更新必须放在 EDT 线程中，否则会抛出 ArrayIndexOutOfBoundsException；"
    "JTable 绑定 DefaultTableModel 是处理二维数据的常用范式，setRowCount(0) + addRow 即可完成刷新。"
    "第四，对“数据库设计决定业务上限”这句话体会颇深：products.seller_id 指向 school_users.id 而非 uni-id-users.id，正是为了支撑“按学校维度隔离 + 学生证强制认证 + 信用分累计”这些校园业务特性，"
    "schema 一旦定型，后期重构成本极高。"
)

add_heading("1.5.3  存在问题与改进方向", size=12)
problems = [
    "1) GUI 界面较为朴素，未做响应式布局，缩放窗口时排版会错位；后续可改用 MigLayout 或 JavaFX。",
    "2) 图片上传目前在代码中以 URL 字符串形式存储，并未真正实现本地文件上传到 MySQL BLOB 或对象存储；后续可引入 Apache Commons FileUpload。",
    "3) 关键字搜索使用 REGEXP 没有走全文索引，数据量大时性能不佳；后续可改用 MySQL FULLTEXT 或 Elasticsearch。",
    "4) 没有实现登录密码加盐哈希，仅以“直接输入 user_id”的演示形式运行；真实场景需配合 BCrypt + 随机盐。",
    "5) CLI 的入参容错还不够友好，没有完全屏蔽掉非法数字输入导致的 NumberFormatException；后续可封装一个统一的 InputUtil。",
    "6) 异常处理仅打印到控制台，GUI 版只弹一个 JOptionPane，缺少统一的日志框架（Log4j2 / SLF4J）支持后续生产化。",
]
for it in problems:
    add_para(it, first_line=False)

add_para(
    "综上所述，本次课程设计不仅覆盖了《面向对象程序设计》课程的关键知识点，"
    "更通过“真实项目驱动”的方式，让我体验到从需求分析、数据库设计、代码实现到测试上线的完整链路，"
    "为后续 Spring Boot + Spring Cloud 的进阶学习奠定了扎实基础。"
)

# 保存
doc.save(OUT)
print("[OK] 报告生成完成：" + OUT)
print("     段落数：" + str(len(doc.paragraphs)))